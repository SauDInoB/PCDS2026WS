#Let's import libs already installed that will be used within the script
import pandas as pd
import matplotlib.pyplot as plt
import os
import json

#Confirm whether they have been installed
print ("importing done")

#You can change the name of the database
filename="asthma_best_db.sav"

#The following block of code will allow for the creation of an dynamic path regardless of the computer the script is running on
current_dir=os.getcwd()
d2s1_folder=os.path.dirname(current_dir)
bd_path="data/raw"

#Now we create a personalized absolute path for the database
db_full_path=os.path.join(d2s1_folder, bd_path,filename)
print (db_full_path)

#We create a new object called my_df. Now we create a personalized absolute path for the database
my_df =pd.read_spss(db_full_path)

# Now we check the shape and columnnames
print("The shape of the dataframe is :", my_df.shape)
columns=my_df.columns.tolist()    
print ("There are", len (columns), "columns in the df")
print("\nColumns are :", columns)

# We check the head and tail of the df
print(my_df.head())
print (my_df.tail())

# We check the datatype of each column
print("\nData types:")
print(my_df.dtypes)

def get_missings(df,column_name):
    num_missing = df[f"{column_name}"].isna().sum()
    result=f"Number of missings: {num_missing}"
    return result

# get missings for every column available in the df
for column in columns:
    column_missings=get_missings(my_df, column)
    print("Missings in", column,":", column_missings)

# # get missings for a specific column available in the df
# column_name="SMQ020"
# print(get_missings(my_df, column_name))

# identificar valores únicos
column_name="SPXBPEF"
unique_values = my_df[f"{column_name}"].unique()
print ("unique values for", column_name, ":", unique_values)

categorical_list=[
    "RIAGENDR",
    "raca_etnia",
    "ENQ100",
    "MCQ010",
    "RDQ031",
    "RDQ050",
    "RDQ070", 
    "RDQ090",
    "RDQ100",
    "RDQ140", 
    "AGQ030", 
    "SMQ020",
    "SMQ040"
]

for i in categorical_list:
    my_df[f"{i}"]= pd.Categorical(my_df[f"{i}"])

print (my_df.dtypes)

# get all values that can be present in a column
# print ("type: ", my_df[f"{column_name}"].dtype) 
# values = my_df[f"{column_name}"].values
# print (values)


from docx import Document

# We can create a new doc to report the findings
doc = Document()

def limpar_doc(path: str):
    """
    Limpa completamente um documento Word existente (.docx):
    - Remove todos os parágrafos
    - Remove todas as tabelas
    - Remove todas as imagens/inline shapes
    Salva de volta no mesmo ficheiro.
    
    Args:
        path (str): Caminho para o ficheiro .docx
    """
    # Abrir documento existente
    doc = Document(path)
    
    # Limpar parágrafos
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # Limpar tabelas
    for table in doc.tables:
        t = table._element
        t.getparent().remove(t)
    
    # Limpar imagens
    for shape in doc.inline_shapes:
        sp = shape._inline
        sp.getparent().remove(sp)
    
    # Salvar documento limpo
    doc.save(path)

doc.add_heading('My summary', level=1)

doc.add_paragraph('This is my report for the Winter School.')
doc.add_paragraph('Columns are' + ';'.join(columns))

for column in columns:
    column_missings=get_missings(my_df, column)
    doc.add_paragraph(column + ":" +column_missings)

doc_name="summary1.docx"
limpar_doc(os.path.join(d2s1_folder, "outputs",doc_name))
doc.save(os.path.join(d2s1_folder, "outputs",doc_name))

# jupytext --to notebook import_and_inspect.py